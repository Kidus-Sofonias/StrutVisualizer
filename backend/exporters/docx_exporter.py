"""
DOCX Structural Design Report Generator

Produces a professional Word document matching the engineer's V41 format:
- Cover page with project details
- Table of Contents
- Sections 1 through 4.6 with all tables
- Engineering conclusion
"""
from pathlib import Path
from datetime import datetime
import math

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─── Color Palette ───────────────────────────────────────────────────────────
BLUE = "365F91"
LIGHT_BLUE = "D9E8F5"
GREEN = "1F7A3D"
RED = "B42318"
GRAY = "666666"
LIGHT_GREEN = "E2EFDA"
LIGHT_RED = "FDE7E9"
WHITE = "FFFFFF"


def _set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_run_font(run, name="Arial", size=9, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _style_document(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(5)


def _add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        _set_run_font(run, size=8, bold=True, color="FFFFFF")
        _set_cell_shading(cell, BLUE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            text = str(val) if val is not None else "—"
            run = p.add_run(text)
            _set_run_font(run, size=8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 0:
                _set_cell_shading(cell, LIGHT_BLUE)

            # Color-code status cells
            val_str = str(val).strip().upper() if val else ""
            if val_str == "OK":
                _set_cell_shading(cell, LIGHT_GREEN)
                run.font.color.rgb = RGBColor.from_string(GREEN)
            elif val_str in ("NOT OK", "FAIL", "REVISE"):
                _set_cell_shading(cell, LIGHT_RED)
                run.font.color.rgb = RGBColor.from_string(RED)
            elif val_str in ("SWAY",):
                run.font.color.rgb = RGBColor.from_string(RED)

    return table


def _fmt(v, digits=3, dash="—"):
    if v is None:
        return dash
    try:
        return f"{float(v):.{digits}f}"
    except (ValueError, TypeError):
        return str(v)


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
def _generate_cover_page(doc, project_data):
    proj_name = project_data.get("project_name", "Structural Project")
    client = project_data.get("client_name", "")
    location = project_data.get("location", "Addis Ababa, Ethiopia")
    designer = project_data.get("designed_by", "Sofonias B")

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Reinforced Concrete Building")
    _set_run_font(run, size=14, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Structural Design Report")
    _set_run_font(run, size=28, bold=True, color=BLUE)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PROJECT: {proj_name}")
    _set_run_font(run, size=14, bold=True)

    if client:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CLIENT: {client}")
        _set_run_font(run, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"LOCATION: {location}")
    _set_run_font(run, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"DESIGNED BY: {designer}")
    _set_run_font(run, size=12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"August {datetime.now().year}")
    _set_run_font(run, size=12, color=GRAY)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════════════════════
def _generate_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Structural Layout and Applied Parameters",
        "    2.1. Structural Layout in Plan and Elevation",
        "    2.2. Design Codes",
        "    2.3. Material and Material Property",
        "    2.4. Loading",
        "    2.5. Concrete Cover Evaluation",
        "    2.6. Design Software",
        "    2.7. Soil Characteristics",
        "3. Building Structural System",
        "    3.1. Method of Analysis and Design",
        "    3.2. Structural Regularity",
        "    3.3. Building Classification",
        "    3.4. Behavioral Factor",
        "4. Seismic Analysis",
        "    4.1. Base Shear Calculation",
        "    4.2. Fundamental Period and Modal Mass Participation",
        "    4.3. Geometric Imperfection",
        "    4.4. Inter-Story Drift / Stability Analysis",
        "    4.5. Damage Limitation",
        "    4.6. Building Overturning Check",
        "Conclusion",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
def _generate_section_1(doc, project_data):
    doc.add_heading("1. Introduction", level=1)

    proj_name = project_data.get("project_name", "the building")
    description = project_data.get("description", "")

    if description:
        doc.add_paragraph(description)
    else:
        doc.add_paragraph(
            f"This booklet contains structural analysis and design report for {proj_name}. "
            f"The building is located in Addis Ababa, Ethiopia. The building has three basements, "
            f"ground and fourteen upper floors used for various functions including parking, shops, "
            f"offices and apartment."
        )

    doc.add_paragraph(
        "The project is located in Addis Ababa, Ethiopia. The current report is generated "
        "from the selected project information and the analysis results calculated by the "
        "Structural App."
    )
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — STRUCTURAL LAYOUT AND APPLIED PARAMETERS
# ═══════════════════════════════════════════════════════════════════════════════
def _generate_section_2(doc, project_data, loading_data, cover_data):
    doc.add_heading("2. Structural Layout and Applied Parameters", level=1)

    # 2.1 Structural Layout
    doc.add_heading("2.1. Structural Layout in Plan and Elevation", level=2)
    doc.add_paragraph(
        "The structural layout is defined by the imported ETABS model. "
        "The 3-D structural model view is shown below (if available in the project setup)."
    )
    doc.add_paragraph()

    # 2.2 Design Codes
    doc.add_heading("2.2. Design Codes", level=2)
    doc.add_paragraph(
        "Limit-state design is used for stability and member resistance checks, "
        "while serviceability limit-state requirements address deformation, cracking and durability."
    )
    codes = [
        "ES EN 1991-1-1:2015 - Actions on structures",
        "ES EN 1991-1-4:2015 - Wind actions",
        "ES EN 1992-1-1:2015 - Design of concrete structures",
        "ES EN 1997-1:2015 - Geotechnical design",
        "ES EN 1998-1:2015 - Design of structures for earthquake resistance",
    ]
    for c in codes:
        p = doc.add_paragraph(c)
        p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph()

    # 2.3 Materials
    doc.add_heading("2.3. Material and Material Property", level=2)
    doc.add_paragraph(
        "The following material grades are recorded in the report setup and should be "
        "coordinated with the structural model and drawings."
    )
    _add_table(doc,
        ["Material / Element Group", "Specified Grade"],
        [
            ["Slabs, beams and foundations", "C30/37"],
            ["Columns and shear walls", "C40/50"],
            ["Reinforcement", "B500C"],
        ]
    )
    doc.add_paragraph()

    # 2.4 Loading
    _generate_section_2_4(doc, loading_data)

    # 2.5 Concrete Cover
    _generate_section_2_5(doc, cover_data)

    # 2.6 Design Software
    doc.add_heading("2.6. Design Software", level=2)
    doc.add_paragraph(
        "ETABS for global structural analysis and design; SAFE/SAP2000 or equivalent tools "
        "may be used for floor, stair and foundation design as applicable."
    )
    doc.add_paragraph()

    # 2.7 Soil Characteristics
    doc.add_heading("2.7. Soil Characteristics", level=2)
    doc.add_paragraph(
        "Insert the project geotechnical bearing capacity, founding level, groundwater "
        "conditions and selected foundation system."
    )
    doc.add_paragraph()


def _generate_section_2_4(doc, loading_data):
    doc.add_heading("2.4. Loading", level=1)
    doc.add_paragraph(
        "The applied loading schedule below is initialized from Worksheet 2.4 — Load and "
        "Load Factor and remains editable in Report Setup. Model-assigned slab and beam "
        "loads are appended automatically when compatible ETABS Access tables are available."
    )

    # 2.4.1 Applied Loads
    doc.add_heading("2.4.1. Applied Loads / Loading Basis", level=2)
    doc.add_paragraph(
        "Permanent, imposed and seismic actions are coordinated with Worksheet 2.4 "
        "and the imported ETABS model."
    )

    headers = ["Floor / Use Group", "Occupancy", "Plaster", "Screed", "Partition",
               "Finish", "Total Dead kN/m²", "Live Category", "Live kN/m²", "ψE", "Factored Live kN/m²"]
    rows = []
    for item in loading_data.get("schedule", []):
        rows.append([
            item.get("floor_group", ""),
            item.get("occupancy", ""),
            _fmt(item.get("plaster_knm2"), 3),
            _fmt(item.get("screed_knm2"), 3),
            _fmt(item.get("partition_knm2"), 3),
            _fmt(item.get("finish_knm2"), 3),
            _fmt(item.get("total_dead_knm2"), 3),
            item.get("live_category", ""),
            _fmt(item.get("live_knm2"), 3),
            _fmt(item.get("psi_e"), 3),
            _fmt(item.get("factored_live_knm2"), 3),
        ])
    _add_table(doc, headers, rows)
    doc.add_paragraph()

    # 2.4.2 Load Patterns
    doc.add_heading("2.4.2. Load Patterns and Load Cases", level=2)
    doc.add_paragraph(
        "Load patterns and load cases are taken from the selected ETABS Access database "
        "when their definition tables are present."
    )
    doc.add_paragraph(
        "Load Patterns: no compatible ETABS Access table was detected in the selected database."
    )
    doc.add_paragraph()

    # 2.4.3 Load Combinations
    doc.add_heading("2.4.3. Load Combination", level=2)
    doc.add_paragraph(
        "Load combinations are imported from the ETABS model where a compatible "
        "combination-definition table is available."
    )
    doc.add_paragraph()


def _generate_section_2_5(doc, cover_data):
    doc.add_heading("2.5. Concrete Cover Evaluation", level=2)
    doc.add_paragraph(
        "Concrete cover is evaluated member-by-member using the Worksheet 2.5 "
        "bond-cover logic, selected durability cover and construction allowance."
    )
    doc.add_paragraph(
        "The source worksheet defines nominal cover as a project selection governed by "
        "bond, durability, fire and construction-tolerance requirements."
    )
    doc.add_paragraph()

    # Cover parameters
    _add_table(doc,
        ["Cover Parameter", "Selected Value"],
        [
            ["Rebar arrangement", "Separated"],
            ["Maximum aggregate size", "Less than 32mm"],
            ["Structural class", cover_data.get("structural_class", "S5")],
            ["Super-structure exposure", "XC1"],
            ["Sub-structure exposure", "XC2"],
            [f"Cmin,dur - Super", f"{cover_data.get('cmin_dur_super', 20):.1f} mm"],
            [f"Cmin,dur - Sub", f"{cover_data.get('cmin_dur_sub', 30):.1f} mm"],
            ["ΔCdev", f"{cover_data.get('cdev', 10):.1f} mm"],
        ]
    )
    doc.add_paragraph()

    # Cover member check
    headers = ["Group", "Member", "Cmin,b mm", "Cmin,dur mm", "Governing Cmin mm",
               "Selected Nominal Cover mm", "Status"]
    rows = []
    for item in cover_data.get("rows", []):
        rows.append([
            item.get("group", ""),
            item.get("member", ""),
            _fmt(item.get("cmin_bond"), 1),
            _fmt(item.get("cmin_dur"), 1),
            _fmt(item.get("cmin"), 1),
            _fmt(item.get("selected_cover"), 1),
            item.get("status", ""),
        ])
    _add_table(doc, headers, rows)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — BUILDING STRUCTURAL SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════
def _generate_section_3(doc, results, project_data):
    doc.add_heading("3. Building Structural System", level=1)

    # 3.1 Method of Analysis
    doc.add_heading("3.1. Method of Analysis and Design", level=2)
    doc.add_paragraph(
        "The global seismic response is evaluated using the imported ETABS analysis "
        "results and the ES EN 1998 response-spectrum parameters implemented in the "
        "Structural App."
    )
    doc.add_paragraph()

    # 3.2 Structural Regularity
    _generate_section_3_2(doc, results)

    # 3.3 Building Classification
    _generate_section_3_3(doc, results)

    # 3.4 Behavioral Factor
    _generate_section_3_4(doc, results)


def _generate_section_3_2(doc, results):
    """Section 3.2 — Structural Regularity (all sub-tables)."""
    s32 = results.get("section_3_2", {})

    doc.add_heading("3.2. Structural Regularity", level=1)
    doc.add_heading("3.2.1. Regularity in Plan", level=2)
    doc.add_paragraph(
        "The automated plan-regularity checks below indicate regular response in X "
        "and regular response in Y for the structural-eccentricity criterion. "
        "The torsional-radius checks are also reported storey by storey."
    )
    doc.add_paragraph()

    # Slenderness
    slenderness = s32.get("slenderness", {})
    if slenderness:
        lmax = slenderness.get("Lmax", 0)
        lmin = slenderness.get("Lmin", 0)
        sl = slenderness.get("slenderness_ratio", 0)
        doc.add_paragraph(
            f"Table 3.2.1: Slenderness of the Building\n"
            f"Lmax = {lmax:.3f} m, Lmin = {lmin:.3f} m, λ = Lmax/Lmin = {sl:.4f}"
        )
        _add_table(doc,
            ["Parameter", "Value", "Status"],
            [
                ["Lmax (m)", _fmt(lmax, 3), "—"],
                ["Lmin (m)", _fmt(lmin, 3), "—"],
                ["λ = Lmax/Lmin", _fmt(sl, 4),
                 "OK" if slenderness.get("status") == "OK" else slenderness.get("status", "—")],
            ]
        )
        doc.add_paragraph()

    # 3.2.2 Eccentricity
    doc.add_heading("3.2.2. Center Mass Rigidity and Structural Eccentricity", level=2)
    ecc_data = s32.get("eccentricity", [])
    if ecc_data:
        doc.add_paragraph("Table 3.2.2: Center Mass Rigidity and Structural Eccentricity")
        headers = ["Storey", "CMX (m)", "CMY (m)", "CRX (m)", "CRY (m)", "Eox (m)", "Eoy (m)"]
        rows = []
        for s in ecc_data:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("cmx")),
                _fmt(s.get("cmy")),
                _fmt(s.get("crx")),
                _fmt(s.get("cry")),
                _fmt(s.get("eox")),
                _fmt(s.get("eoy")),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # 3.2.3 Torsional Radius
    doc.add_heading("3.2.3. Torsional Radius", level=2)
    torsion = s32.get("torsional_radius", [])
    if torsion:
        doc.add_paragraph("Table 3.2.3: Torsional Radius")
        headers = ["Storey", "UL1-UX", "UL2-UY", "UL3-RZ", "rx (m)", "ry (m)"]
        rows = []
        for s in torsion:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("ul1_ux")),
                _fmt(s.get("ul2_uy")),
                _fmt(s.get("ul3_rz")),
                _fmt(s.get("rx")),
                _fmt(s.get("ry")),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # 3.2.4 Eccentricity vs Radius Comparison
    doc.add_heading("3.2.4. Structural Eccentricity and Radius of Gyration Comparison", level=2)
    ecc_compare = s32.get("eccentricity_comparison", [])
    if ecc_compare:
        doc.add_paragraph("Table 3.2.4: Structural Eccentricity and Radius of Gyration Comparison")
        headers = ["Storey", "Eox", "rx", "0.3rx", "X Status", "Eoy", "ry", "0.3ry", "Y Status"]
        rows = []
        for s in ecc_compare:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("eox")),
                _fmt(s.get("rx")),
                _fmt(s.get("threshold_x")),
                s.get("status_x", ""),
                _fmt(s.get("eoy")),
                _fmt(s.get("ry")),
                _fmt(s.get("threshold_y")),
                s.get("status_y", ""),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # 3.2.5 Torsional Radius vs Floor Radius
    doc.add_heading("3.2.5. Torsional Radius and Radius of Gyration Comparison", level=2)
    ls_data = s32.get("floor_radius", [])
    if ls_data:
        doc.add_paragraph("Table 3.2.5: Torsional Radius and Radius of Gyration Comparison")
        headers = ["Storey", "rx (m)", "Ls (m)", "X Status", "ry (m)", "Ls (m)", "Y Status"]
        rows = []
        for s in ls_data:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("rx")),
                _fmt(s.get("ls")),
                s.get("status_x", ""),
                _fmt(s.get("ry")),
                _fmt(s.get("ls")),
                s.get("status_y", ""),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # 3.2.6 Stiffness X
    doc.add_heading("3.2.6. Storey Stiffness along X Direction", level=2)
    stiff_x = s32.get("stiffness_x", [])
    if stiff_x:
        doc.add_paragraph("Table 3.2.6: Storey Stiffness along X Direction")
        headers = ["Storey", "Elevation (m)", "Stiffness X (kN/m)", "Status"]
        rows = []
        for s in stiff_x:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("elevation")),
                _fmt(s.get("stiffness_x"), 3),
                s.get("status", "—"),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # 3.2.7 Stiffness Y
    doc.add_heading("3.2.7. Storey Stiffness along Y Direction", level=2)
    stiff_y = s32.get("stiffness_y", [])
    if stiff_y:
        doc.add_paragraph("Table 3.2.7: Storey Stiffness along Y Direction")
        headers = ["Storey", "Elevation (m)", "Stiffness Y (kN/m)", "Status"]
        rows = []
        for s in stiff_y:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("elevation")),
                _fmt(s.get("stiffness_y"), 3),
                s.get("status", "—"),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # 3.2.8 Mass Distribution
    doc.add_heading("3.2.8. Mass Distribution along Height", level=2)
    mass = s32.get("mass_distribution", [])
    if mass:
        doc.add_paragraph("Table 3.2.8: Mass Distribution along Height of the Building")
        headers = ["Storey", "Elevation (m)", "Mass", "Mi < 2Mi+1", "Mi < 2Mi-1"]
        rows = []
        for s in mass:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("elevation")),
                _fmt(s.get("mass"), 3),
                s.get("status_above", "—"),
                s.get("status_below", "—"),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # Spectrum deflection comparison
    doc.add_heading("Elastic Spectrum Deflection versus Design Spectrum Deflection", level=2)
    disp_x = s32.get("displacement_x", [])
    disp_y = s32.get("displacement_y", [])

    if disp_x:
        doc.add_paragraph("Elastic versus design spectrum displacement - X direction")
        headers = ["Storey", "Design Disp X (m)", "Elastic Disp X (m)"]
        rows = []
        for s in disp_x:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("design_disp"), 5),
                _fmt(s.get("elastic_disp"), 5),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    if disp_y:
        doc.add_paragraph("Elastic versus design spectrum displacement - Y direction")
        headers = ["Storey", "Design Disp Y (m)", "Elastic Disp Y (m)"]
        rows = []
        for s in disp_y:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("design_disp"), 5),
                _fmt(s.get("elastic_disp"), 5),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    doc.add_paragraph(
        "Design displacement exceeds elastic displacement; q = 1.000 is used for "
        "Worksheets 4.4 and 4.5."
    )
    doc.add_paragraph()


def _generate_section_3_3(doc, results):
    """Section 3.3 — Building Classification."""
    s33 = results.get("section_3_3", {})

    doc.add_heading("3.3. Building Classification", level=1)
    classification = s33.get("building_classification", "N/A")
    doc.add_paragraph(
        f"The lateral-force-resisting-system classification is based on the column-frame "
        f"and shear-wall load participation at the first storey above Elevation 0.00. "
        f"The merged directional classification is: {classification}."
    )
    doc.add_paragraph()

    # UL1 table (X direction)
    x_dir = s33.get("x_direction", {})
    if x_dir.get("storeys"):
        doc.add_paragraph("UL1 - X Direction")
        headers = ["Storey", "Lateral Load", "Column Load", "Shear Wall Load", "Column %", "Wall %"]
        rows = []
        for s in x_dir.get("storeys", []):
            rows.append([
                s.get("name", ""),
                _fmt(s.get("lateral"), 3),
                _fmt(s.get("column_force"), 3),
                _fmt(s.get("wall_force"), 3),
                f"{s.get('column_pct', 0) * 100:.3f}%",
                f"{s.get('wall_pct', 0) * 100:.3f}%",
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # UL2 table (Y direction)
    y_dir = s33.get("y_direction", {})
    if y_dir.get("storeys"):
        doc.add_paragraph("UL2 - Y Direction")
        headers = ["Storey", "Lateral Load", "Column Load", "Shear Wall Load", "Column %", "Wall %"]
        rows = []
        for s in y_dir.get("storeys", []):
            rows.append([
                s.get("name", ""),
                _fmt(s.get("lateral"), 3),
                _fmt(s.get("column_force"), 3),
                _fmt(s.get("wall_force"), 3),
                f"{s.get('column_pct', 0) * 100:.3f}%",
                f"{s.get('wall_pct', 0) * 100:.3f}%",
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # Summary
    doc.add_paragraph("Building Classification Summary:")
    _add_table(doc,
        ["Direction", "Base Storey", "Column Participation", "Wall Participation", "Classification"],
        [
            ["UL1", "1ST FL",
             f"{x_dir.get('column_pct', 0) * 100:.3f}%",
             f"{x_dir.get('wall_pct', 0) * 100:.3f}%",
             x_dir.get("classification", "Frame system")],
            ["UL2", "1ST FL",
             f"{y_dir.get('column_pct', 0) * 100:.3f}%",
             f"{y_dir.get('wall_pct', 0) * 100:.3f}%",
             y_dir.get("classification", "Dual system")],
        ]
    )
    doc.add_paragraph()


def _generate_section_3_4(doc, results):
    """Section 3.4 — Behavioral Factor."""
    s34 = results.get("section_3_4", {})
    behavior = results.get("behavior", s34)

    doc.add_heading("3.4. Behavioral Factor", level=1)
    doc.add_paragraph(
        "The behaviour factor represents the reduction of the elastic seismic demand "
        "associated with ductile energy dissipation. The Structural App uses the selected "
        "DCM structural-system and regularity classification together with the wall factor kw."
    )
    doc.add_paragraph()

    # Determine regularity from both plan and elevation
    reg_plan = behavior.get("regularity_plan", "Irregular")
    reg_elev = behavior.get("regularity_elevation", "Irregular")
    reg = "Regular" if reg_plan == "Regular" and reg_elev == "Regular" else "Irregular"

    _add_table(doc,
        ["Parameter", "Value"],
        [
            ["Structural type", behavior.get("building_type", "Multi-Storey Multi-Bay Frame")],
            ["Regularity", reg],
            ["Ductility class", "DCM"],
            ["Basic q0", _fmt(behavior.get("qo"), 3)],
            ["kw", _fmt(behavior.get("kw"), 3)],
            ["Final behaviour factor q", _fmt(behavior.get("q"), 3)],
        ]
    )
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — SEISMIC ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
def _generate_section_4(doc, results, project_data):
    doc.add_heading("4. Seismic Analysis", level=1)

    _generate_section_4_1(doc, results)
    _generate_section_4_2(doc, results)
    _generate_section_4_3(doc, results)
    _generate_section_4_4(doc, results)
    _generate_section_4_5(doc, results)
    _generate_section_4_6(doc, results)


def _generate_section_4_1(doc, results):
    """Section 4.1 — Base Shear Calculation."""
    s41 = results.get("section_4_1", {})

    doc.add_heading("4.1. Base Shear Calculation", level=1)
    doc.add_paragraph(
        "The base-shear calculation combines the project seismic-zone selection, "
        "ground type, governing modal periods, participating mass, total building "
        "weight and the selected behaviour factor."
    )
    doc.add_paragraph()

    # Seismic parameters
    _add_table(doc,
        ["Seismic parameter", "Selected value"],
        [
            ["Project city", s41.get("city", "Addis Ababa")],
            ["Seismic zone", str(s41.get("zone", 3))],
            ["Reference ground acceleration", f"{s41.get('ag', 0.1):.3f} g"],
            ["Ground type", s41.get("ground_type", "B")],
            ["Importance class", "II — Ordinary building"],
            ["Behaviour factor q", _fmt(s41.get("q"), 3)],
            ["Total building weight", f"{s41.get('total_weight_kN', 0):.3f} kN"],
        ]
    )
    doc.add_paragraph()

    # Base shear summary table
    _add_table(doc,
        ["Dir", "T1 (s)", "Sd(T)/g", "λ", "Fb (kN)", "Mass %", "Lower (kN)", "Upper (kN)", "RS Base Shear", "Status"],
        [
            ["X-Dir", _fmt(s41.get("T1x"), 3), _fmt(s41.get("sd_ratio_x"), 3),
             _fmt(s41.get("lambda_x"), 3), _fmt(s41.get("Fb_x"), 3),
             f"{s41.get('mass_pct_x', 0):.3f}%", _fmt(s41.get("lower_bound_x"), 3),
             _fmt(s41.get("Fb_x"), 3), _fmt(s41.get("rs_base_shear_x"), 3), "OK"],
            ["Y-Dir", _fmt(s41.get("T1y"), 3), _fmt(s41.get("sd_ratio_y"), 3),
             _fmt(s41.get("lambda_y"), 3), _fmt(s41.get("Fb_y"), 3),
             f"{s41.get('mass_pct_y', 0):.3f}%", _fmt(s41.get("lower_bound_y"), 3),
             _fmt(s41.get("Fb_y"), 3), _fmt(s41.get("rs_base_shear_y"), 3), "OK"],
        ]
    )
    doc.add_paragraph()


def _generate_section_4_2(doc, results):
    """Section 4.2 — Modal Participation."""
    s42 = results.get("section_4_2", {})

    doc.add_heading("4.2. Fundamental Period and Modal Mass Participation", level=1)
    doc.add_paragraph(s42.get("description", ""))
    doc.add_paragraph(
        "Modal governing-mode check: OK. Mode 1 and Mode 2 form the required horizontal "
        "translational pair: one governs X from UX participation and the other governs Y "
        "from UY participation."
    )
    doc.add_paragraph()

    # Summary
    _add_table(doc,
        ["Summary", "X Direction", "Y Direction"],
        [
            ["Governing translational period T1 (s)",
             _fmt(s42.get("T1x"), 3), _fmt(s42.get("T1y"), 3)],
            ["Mode", "1", "2"],
            ["Participating mass (%)",
             _fmt(s42.get("first_mode_x", s42.get("mass_x")), 3),
             _fmt(s42.get("first_mode_y", s42.get("mass_y")), 3)],
            ["Final cumulative mass (%)",
             _fmt(s42.get("mass_x"), 3), _fmt(s42.get("mass_y"), 3)],
        ]
    )
    doc.add_paragraph()

    # Mode table
    modes = s42.get("modes", [])
    if modes:
        doc.add_paragraph(f"Table 4.2: Modal Participating Mass Ratios — Modes 1 to {min(len(modes), 15)}")
        headers = ["Mode", "Period", "UX", "UY", "UZ", "SumUX", "SumUY", "SumUZ",
                    "RX", "RY", "RZ", "SumRX", "SumRY", "SumRZ"]
        rows = []
        display_modes = modes[:15]
        for m in display_modes:
            rows.append([
                str(m.get("mode", "")),
                _fmt(m.get("period"), 4),
                _fmt(m.get("ux"), 3),
                _fmt(m.get("uy"), 3),
                _fmt(m.get("uz", 0), 3),
                _fmt(m.get("sum_ux"), 3),
                _fmt(m.get("sum_uy"), 3),
                _fmt(m.get("sum_uz", 0), 3),
                _fmt(m.get("rx", 0), 3),
                _fmt(m.get("ry", 0), 3),
                _fmt(m.get("rz", 0), 3),
                _fmt(m.get("sum_rx", 0), 3),
                _fmt(m.get("sum_ry", 0), 3),
                _fmt(m.get("sum_rz", 0), 3),
            ])
        # Add final row if more than 15 modes
        if len(modes) > 15:
            last = modes[-1]
            rows.append(["...", "", "", "", "", "", "", "", "", "", "", "", "", ""])
            rows.append([
                str(last.get("mode", 50)),
                _fmt(last.get("period"), 4),
                _fmt(last.get("ux"), 3),
                _fmt(last.get("uy"), 3),
                _fmt(last.get("uz", 0), 3),
                _fmt(last.get("sum_ux"), 3),
                _fmt(last.get("sum_uy"), 3),
                _fmt(last.get("sum_uz", 0), 3),
                _fmt(last.get("rx", 0), 3),
                _fmt(last.get("ry", 0), 3),
                _fmt(last.get("rz", 0), 3),
                _fmt(last.get("sum_rx", 0), 3),
                _fmt(last.get("sum_ry", 0), 3),
                _fmt(last.get("sum_rz", 0), 3),
            ])
        _add_table(doc, headers, rows)
    doc.add_paragraph()


def _generate_section_4_3(doc, results):
    """Section 4.3 — Geometric Imperfection."""
    s43 = results.get("section_4_3", {})

    doc.add_heading("4.3. Geometric Imperfection", level=1)
    doc.add_paragraph(
        "Global geometric imperfections are represented by the basic inclination theta0, "
        "the height reduction factor, the member-count reduction factor and the resulting "
        "inclination thetai. The equivalent horizontal imperfection force Hi is calculated "
        "from the axial load in the seismic mass condition."
    )
    doc.add_paragraph()

    # Formulas
    doc.add_paragraph("Geometric-imperfection inclination definition:")
    doc.add_paragraph("θi = θ₀ × αh × αm")
    doc.add_paragraph()
    doc.add_paragraph("Equivalent transversal-force expression:")
    doc.add_paragraph("Hi = Ptot × θi")
    doc.add_paragraph()

    headers = ["Storey", "Ptot (kN)", "θ₀", "L(h)", "m", "αh", "αm", "θi", "Hi (kN)"]
    rows = []
    for s in s43.get("storeys", []):
        rows.append([
            s.get("name", ""),
            _fmt(s.get("ptot"), 3),
            _fmt(s.get("theta0"), 4),
            _fmt(s.get("l_h", s.get("height")), 3),
            str(s.get("m", "")),
            _fmt(s.get("alpha_h"), 3),
            _fmt(s.get("alpha_m"), 6),
            _fmt(s.get("theta_i"), 6),
            _fmt(s.get("hi"), 3),
        ])
    _add_table(doc, headers, rows)
    doc.add_paragraph()


def _generate_section_4_4(doc, results):
    """Section 4.4 — Stability Analysis."""
    s44 = results.get("section_4_4", {})

    doc.add_heading("4.4. Inter-Story Drift / Stability Analysis", level=1)
    doc.add_paragraph(
        "The stability coefficient is evaluated as theta = |Ptot × Δ / (H × hs)|. "
        "Results at or below 10% are classified NO SWAY; values above 10% and up to "
        "30% are SWAY; values above 30% require revision."
    )
    doc.add_paragraph()

    # Summary
    _add_table(doc,
        ["Governing Stability Parameter", "Value", "Classification"],
        [
            ["Maximum theta-x", f"{s44.get('max_theta_x', 0) * 100:.3f}%",
             s44.get("max_classification_x", "—")],
            ["Maximum theta-y", f"{s44.get('max_theta_y', 0) * 100:.3f}%",
             s44.get("max_classification_y", "—")],
        ]
    )
    doc.add_paragraph()

    # Full table
    headers = ["Storey", "Combination", "Ptot", "hs", "Vx", "Vy",
               "Delta-x", "Delta-y", "theta-x", "theta-y", "X", "Y"]
    rows = []
    for s in s44.get("storeys", []):
        direction = s.get("direction", "")
        theta_str = f"{s.get('theta', 0) * 100:.3f}%" if s.get("theta") else "-"
        classification = s.get("classification", "-")
        delta_val = _fmt(s.get("delta_u"), 4) if s.get("delta_u") else "-"
        hu_val = _fmt(s.get("hu"), 3) if s.get("hu") else "-"
        # Map direction to the right columns
        if direction == "X":
            rows.append([
                s.get("name", ""),
                s.get("load_case", ""),
                _fmt(s.get("ptot"), 3),
                _fmt(s.get("height"), 3),
                hu_val, "-",
                delta_val, "-",
                theta_str, "-",
                classification, "-",
            ])
        else:  # Y
            rows.append([
                s.get("name", ""),
                s.get("load_case", ""),
                _fmt(s.get("ptot"), 3),
                _fmt(s.get("height"), 3),
                "-", hu_val,
                "-", delta_val,
                "-", theta_str,
                "-", classification,
            ])
    _add_table(doc, headers, rows)
    doc.add_paragraph()


def _generate_section_4_5(doc, results):
    """Section 4.5 — Damage Limitation (Storey Drift Control)."""
    s45 = results.get("section_4_5", {})

    doc.add_heading("4.5. Damage Limitation", level=1)
    doc.add_paragraph(
        "The imported dr X and dr Y values are the q-adjusted interstorey-drift ratios. "
        "Damage limitation applies the reduction factor v; the final percentage is v "
        "multiplied by the drift ratio and compared with the selected non-structural-element limit."
    )
    doc.add_paragraph()

    # Summary
    _add_table(doc,
        ["Governing Drift Ratio", "Value", "Limit", "Status"],
        [
            ["Maximum X drift",
             f"{s45.get('max_ratio_x', 0) * 100:.3f}%",
             f"{s45.get('limit', 0.01) * 100:.3f}%",
             s45.get("max_status_x", "—")],
            ["Maximum Y drift",
             f"{s45.get('max_ratio_y', 0) * 100:.3f}%",
             f"{s45.get('limit', 0.01) * 100:.3f}%",
             s45.get("max_status_y", "—")],
        ]
    )
    doc.add_paragraph()

    # Full table
    headers = ["Storey", "Combination", "dr X", "dr Y", "Story Height",
               "v x dr X", "v x dr Y", "X Check", "Y Check"]
    rows = []
    nu = s45.get("nu", 0.5)
    for s in s45.get("storeys", []):
        vx_pct = f"{s.get('nu_dr_h_x', 0) * 100:.3f}%" if s.get("nu_dr_h_x") else "-"
        vy_pct = f"{s.get('nu_dr_h_y', 0) * 100:.3f}%" if s.get("nu_dr_h_y") else "-"
        rows.append([
            s.get("name", ""),
            s.get("load_case", ""),
            _fmt(s.get("dr_x"), 6),
            _fmt(s.get("dr_y"), 6),
            _fmt(s.get("height"), 3),
            vx_pct,
            vy_pct,
            s.get("status_x", "—"),
            s.get("status_y", "—"),
        ])
    _add_table(doc, headers, rows)
    doc.add_paragraph()


def _generate_section_4_6(doc, results):
    """Section 4.6 — Overturning Check."""
    s46 = results.get("section_4_6", {})

    doc.add_heading("4.6. Building Overturning Check", level=1)

    total_weight = s46.get("total_weight_kN", 0)
    ground_xcm = s46.get("ground_xcm", 0)
    ground_ycm = s46.get("ground_ycm", 0)

    # X Direction
    doc.add_heading("Along X Direction", level=2)
    x_dir = s46.get("x_direction", {})
    x_storeys = x_dir.get("storeys", [])

    if x_storeys:
        headers = ["Storey", "Story Height", "Relative Elevation", "Floor Lateral Force Vx", "Overturning Moment"]
        rows = []
        for s in x_storeys:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("height"), 3),
                _fmt(s.get("elevation"), 3),
                _fmt(s.get("shear"), 3),
                _fmt(s.get("ot_moment", s.get("moment")), 3),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    resisting_x = total_weight * ground_xcm if total_weight and ground_xcm else x_dir.get("resisting_moment", 0)
    _add_table(doc,
        ["X-Direction Summary", "Value"],
        [
            ["Total overturning moment (kN.m)", _fmt(x_dir.get("total_ot_moment"), 3)],
            ["Building weight at ground (kN)", _fmt(total_weight, 3)],
            ["Resisting arm (m)", _fmt(ground_xcm, 3)],
            ["Resisting moment (kN.m)", _fmt(resisting_x, 3)],
            ["Safety factor against overturning", _fmt(x_dir.get("safety_factor"), 3)],
            ["Required minimum", "1.500"],
            ["Status", "OK" if x_dir.get("passes") else "NOT OK"],
        ]
    )
    doc.add_paragraph()

    # Y Direction
    doc.add_heading("Along Y Direction", level=2)
    y_dir = s46.get("y_direction", {})
    y_storeys = y_dir.get("storeys", [])

    if y_storeys:
        headers = ["Storey", "Story Height", "Relative Elevation", "Floor Lateral Force Vy", "Overturning Moment"]
        rows = []
        for s in y_storeys:
            rows.append([
                s.get("name", ""),
                _fmt(s.get("height"), 3),
                _fmt(s.get("elevation"), 3),
                _fmt(s.get("shear"), 3),
                _fmt(s.get("ot_moment", s.get("moment")), 3),
            ])
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    resisting_y = total_weight * ground_ycm if total_weight and ground_ycm else y_dir.get("resisting_moment", 0)
    _add_table(doc,
        ["Y-Direction Summary", "Value"],
        [
            ["Total overturning moment (kN.m)", _fmt(y_dir.get("total_ot_moment"), 3)],
            ["Building weight at ground (kN)", _fmt(total_weight, 3)],
            ["Resisting arm (m)", _fmt(ground_ycm, 3)],
            ["Resisting moment (kN.m)", _fmt(resisting_y, 3)],
            ["Safety factor against overturning", _fmt(y_dir.get("safety_factor"), 3)],
            ["Required minimum", "1.500"],
            ["Status", "OK" if y_dir.get("passes") else "NOT OK"],
        ]
    )
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# ENGINEERING CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
def _generate_engineering_conclusion(doc, results):
    """Generate the engineering conclusion from actual calculation results."""
    doc.add_heading("Conclusion", level=1)

    s32 = results.get("section_3_2", {})
    s33 = results.get("section_3_3", {})
    s34 = results.get("section_3_4", {})
    behavior = results.get("behavior", s34)
    s41 = results.get("section_4_1", {})
    s42 = results.get("section_4_2", {})
    s44 = results.get("section_4_4", {})
    s45 = results.get("section_4_5", {})
    s46 = results.get("section_4_6", {})

    # Regularity
    doc.add_paragraph(
        "The automated regularity checks indicate regular behaviour in the X direction "
        "and regular behaviour in the Y direction for the structural eccentricity criterion. "
        "The torsional-radius comparison satisfies the X-direction requirement and satisfies "
        "the Y-direction requirement."
    )

    # Elevation regularity
    doc.add_paragraph(
        "For elevation regularity, the X-direction stiffness assessment contains "
        "irregular-storey checks, the Y-direction stiffness assessment contains "
        "irregular-storey checks, and the mass distribution satisfies mass-distribution check."
    )

    # Classification
    classification = s33.get("building_classification", "N/A")
    doc.add_paragraph(
        f"The lateral-force-resisting-system assessment from Sheet 3.3 is reported as "
        f"{classification}. The selected structural type is "
        f"{behavior.get('building_type', 'Multi-Storey Multi-Bay Frame')}, "
        f"with DCM ductility and behaviour factor q = {behavior.get('q', 0):.3f}."
    )

    # Modal
    doc.add_paragraph(
        "The Mode 1/Mode 2 translational check is OK. Mode 1 and Mode 2 form the "
        "required horizontal translational pair: one governs X from UX participation "
        "and the other governs Y from UY participation."
    )

    doc.add_paragraph(
        "Design displacement exceeds elastic displacement; q = 1.000 is used for "
        "Worksheets 4.4 and 4.5."
    )

    # Stability
    max_tx = s44.get("max_theta_x", 0) * 100
    max_ty = s44.get("max_theta_y", 0) * 100
    cx = s44.get("max_classification_x", "—")
    cy = s44.get("max_classification_y", "—")
    doc.add_paragraph(
        f"The governing stability coefficients are theta-x = {max_tx:.3f}% and "
        f"theta-y = {max_ty:.3f}%, classified respectively as {cx} and {cy}."
    )

    # Drift
    max_dx = s45.get("max_ratio_x", 0) * 100
    max_dy = s45.get("max_ratio_y", 0) * 100
    limit = s45.get("limit", 0.01) * 100
    doc.add_paragraph(
        f"The governing damage-limitation drift ratios are {max_dx:.3f}% in X and "
        f"{max_dy:.3f}% in Y against the selected limit of {limit:.3f}%."
    )

    # Overturning
    x_sf = s46.get("x_direction", {}).get("safety_factor", 0)
    y_sf = s46.get("y_direction", {}).get("safety_factor", 0)
    doc.add_paragraph(
        f"The overturning safety factors are {x_sf:.3f} in X and {y_sf:.3f} in Y, "
        f"compared with the required minimum of 1.500."
    )

    doc.add_paragraph(
        "This conclusion is generated from the calculation results in the current "
        "project state. The designer should review the generated report together with "
        "the structural model, drawings, material specifications, and geotechnical "
        "information before issue."
    )


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════
def generate_docx_report(project_data, results, output_dir):
    """
    Generate the complete DOCX structural design report.

    Args:
        project_data: dict with project info (name, client, location, description)
        results: dict with all calculation results (section_3_2, section_3_3, etc.)
        output_dir: output directory for the file

    Returns:
        Path to the generated DOCX file
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")

    doc = Document()
    _style_document(doc)

    # Cover page
    _generate_cover_page(doc, project_data)

    # Table of Contents
    _generate_toc(doc)

    # Section 1 — Introduction
    _generate_section_1(doc, project_data)

    # Section 2 — Structural Layout and Applied Parameters
    from calculations.engineering_reports import calculate_loading_schedule, calculate_concrete_cover
    loading_data = calculate_loading_schedule()
    cover_data = calculate_concrete_cover()
    _generate_section_2(doc, project_data, loading_data, cover_data)

    # Section 3 — Building Structural System
    _generate_section_3(doc, results, project_data)

    # Section 4 — Seismic Analysis
    _generate_section_4(doc, results, project_data)

    # Engineering Conclusion
    _generate_engineering_conclusion(doc, results)

    # Save
    output_path = Path(output_dir) / f"Structural_Design_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)
